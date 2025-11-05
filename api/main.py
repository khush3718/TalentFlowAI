import os
import re
import io
import json
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from string import Template

import pdfplumber
import docx
from docx import Document
import google.generativeai as genai
from fastapi import FastAPI, File, UploadFile, HTTPException, status, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from fastapi import Body
# ==================== CONFIGURATION ====================
load_dotenv()

# Current date for experience calculations
CURRENT_DATE = datetime.now()
CURRENT_YEAR = CURRENT_DATE.year
CURRENT_MONTH = CURRENT_DATE.month
CURRENT_DATE_STR = CURRENT_DATE.strftime("%B %Y")  # e.g., "November 2025"

# Configure logging for production
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]  # Only StreamHandler for Render
)
logger = logging.getLogger(__name__)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# GOOGLE_API_KEY = ""

if not GOOGLE_API_KEY:
    logger.error("Missing GOOGLE_API_KEY in environment variables")
    raise ValueError("GOOGLE_API_KEY must be set in environment variables")

genai.configure(api_key=GOOGLE_API_KEY)

# ==================== PYDANTIC MODELS ====================
class WorkExperience(BaseModel):
    company: str
    role: Optional[str] = None
    duration: str  # e.g., "2 years 3 months" or "Jan 2020 - Dec 2022"
    years: float  # Numeric years (e.g., 2.25)


class Education(BaseModel):
    degree: str
    institution: Optional[str] = None
    year: Optional[str] = None
    field_of_study: Optional[str] = None


class CandidateDetails(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    work_experience: List[WorkExperience] = Field(default_factory=list)
    education: List[Education] = Field(default_factory=list)
    total_experience_years: Optional[float] = None


class KeywordItem(BaseModel):
    keyword: str
    relevance_score: int = Field(..., ge=0, le=100)


class Keywords(BaseModel):
    keywords: List[KeywordItem] = Field(default_factory=list, max_length=8)


class ResumeData(BaseModel):
    job_role: str
    job_description: str
    candidate_details: CandidateDetails
    keywords: Keywords
    extracted_at: str = Field(default_factory=lambda: datetime.now().isoformat())


class Suggestion(BaseModel):
    missing_skill: str
    suggested_action: str
    priority: str = "medium"


class ScreeningAnalysis(BaseModel):
    summary: str
    ats_score: int = Field(..., ge=0, le=100)
    grade: str
    suggestions: List[Suggestion]
    analyzed_at: str = Field(default_factory=lambda: datetime.now().isoformat())


class CompleteAnalysisResponse(BaseModel):
    success: bool
    resume_data: ResumeData
    screening: ScreeningAnalysis
    analyzed_at: str


class HealthCheck(BaseModel):
    status: str
    message: str
    timestamp: str
    version: str = "2.3.0"


class ReportRequest(BaseModel):
    analysis_data: CompleteAnalysisResponse


class ErrorResponse(BaseModel):
    error: str
    detail: str
    timestamp: str = Field(default_factory=lambda: datetime.now().isoformat())

# ==================== SERVICES ====================
class DocumentExtractor:
    @staticmethod
    def extract_text_from_pdf(pdf_file: bytes) -> str:
        """Extract raw text from PDF"""
        try:
            text = ''
            with pdfplumber.open(io.BytesIO(pdf_file)) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ''
                    text += page_text + '\n'
                    logger.info(f"Extracted text from PDF page {i}")
            if not text.strip():
                raise ValueError("No text found in PDF.")
            logger.info(f"Total extracted characters from PDF: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise HTTPException(status_code=500, detail=f"PDF extraction failed: {e}")
    
    @staticmethod
    def extract_text_from_docx(docx_file: bytes) -> str:
        """Extract raw text from DOCX"""
        try:
            doc = Document(io.BytesIO(docx_file))
            text = ''
            
            # Extract text from paragraphs
            for i, paragraph in enumerate(doc.paragraphs, start=1):
                text += paragraph.text + '\n'
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + '\n'
            
            if not text.strip():
                raise ValueError("No text found in DOCX.")
            
            logger.info(f"Total extracted characters from DOCX: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise HTTPException(status_code=500, detail=f"DOCX extraction failed: {e}")
    
    @staticmethod
    def extract_text_from_doc(doc_file: bytes) -> str:
        """Extract raw text from DOC (older format)"""
        try:
            # Try to read as DOCX first (some .doc files are actually .docx)
            try:
                return DocumentExtractor.extract_text_from_docx(doc_file)
            except:
                # If that fails, we need python-docx2txt or antiword
                # For now, we'll raise an error with a helpful message
                raise HTTPException(
                    status_code=400, 
                    detail="Legacy .doc format detected. Please convert to .docx or .pdf format."
                )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOC extraction error: {e}")
            raise HTTPException(status_code=500, detail=f"DOC extraction failed: {e}")
    
    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        """Extract text based on file extension"""
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return DocumentExtractor.extract_text_from_pdf(file_bytes)
        elif file_extension == 'docx':
            return DocumentExtractor.extract_text_from_docx(file_bytes)
        elif file_extension == 'doc':
            return DocumentExtractor.extract_text_from_doc(file_bytes)
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file format: .{file_extension}. Supported formats: .pdf, .docx, .doc"
            )


class TextCleaner:
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        text = re.sub(r'\(cid:\d+\)', '', text)
        text = ''.join(ch for ch in text if ord(ch) >= 32 or ch in '\n\t\r')
        lines = text.split('\n')
        cleaned = [re.sub(r' +', ' ', ln.strip()) for ln in lines if ln.strip()]
        return '\n'.join(cleaned)


class AIService:
    """AI analysis service using Google Gemini"""

    def __init__(self):
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        logger.info("Gemini model initialized for ATS analysis")

    # ---------- PROMPT 1: Candidate Data Extraction ----------
    def extract_candidate_data(self, cv_text: str) -> Dict[str, Any]:
        prompt_tmpl = Template("""
You are an AI assistant specialized in parsing resumes. Extract the following details comprehensively.

CURRENT DATE: ${current_date} (Use this to calculate durations for ongoing positions)
CURRENT YEAR: ${current_year}

RESUME TEXT:
${cv_text}

OUTPUT JSON FORMAT:
{
  "candidate_details": {
    "name": "<full name>",
    "email": "<email>",
    "phone": "<phone>",
    "location": "<location>",
    "total_experience_years": <total years as decimal, e.g., 5.5>,
    "work_experience": [
      {
        "company": "<company name>",
        "role": "<job title>",
        "duration": "<e.g., Jan 2020 - Dec 2022 or 2 years 3 months>",
        "years": <numeric years, e.g., 2.25>
      }
    ],
    "education": [
      {
        "degree": "<degree name, e.g., Bachelor of Science>",
        "institution": "<university/college name>",
        "year": "<graduation year or date range>",
        "field_of_study": "<major/specialization>"
      }
    ]
  },
  "keywords": {
    "keywords": [
      {"keyword": "Python", "relevance_score": 95},
      {"keyword": "AWS", "relevance_score": 88}
    ]
  }
}

EXTRACTION RULES:
1. **Work Experience**: Extract ALL companies mentioned with their respective durations
   - Calculate years as decimal (e.g., 1 year 6 months = 1.5 years)
   - If only dates given, calculate the duration from start date to end date
   - For ongoing positions (e.g., "2024 - Present", "Jan 2023 - Current"), calculate duration up to ${current_date}
   - Convert all date ranges to numeric years with 2 decimal places
   - Include current positions (use "Present" for end date in duration field)
   
2. **Total Experience**: Sum up all work experience years (calculate ongoing positions up to current date)

3. **Education**: Extract ALL educational qualifications
   - Include degree type, institution, graduation year, and field of study
   - List from highest to lowest degree
   
4. **Keywords**: Always return exactly 8 most relevant keywords with scores (0-100)

5. Return ONLY valid JSON, no additional text.
""")
        try:
            prompt = prompt_tmpl.substitute(
                cv_text=cv_text,
                current_date=CURRENT_DATE_STR,
                current_year=CURRENT_YEAR
            )
            response = self.model.generate_content(prompt)
            
            # Handle response properly
            if hasattr(response, 'text'):
                result_text = response.text.strip()
            elif hasattr(response, 'parts'):
                result_text = ''.join(part.text for part in response.parts if hasattr(part, 'text')).strip()
            else:
                result_text = str(response).strip()

            # Extract JSON from markdown code blocks
            json_match = re.search(r'```json\s*(.*?)\s*```', result_text, re.DOTALL | re.IGNORECASE)
            if json_match:
                result_text = json_match.group(1).strip()
            
            # Parse JSON
            data = json.loads(result_text)
            
            # Validate structure
            if not isinstance(data, dict):
                raise ValueError("Response is not a valid JSON object")
            
            return data
        except Exception as e:
            return {
                "candidate_details": {
                    "work_experience": [],
                    "education": []
                },
                "keywords": {"keywords": []}
            }

    # ---------- PROMPT 2: Comprehensive ATS Analysis ----------
    def analyze_resume_comprehensive(
        self, job_role: str, job_description: str, cv_text: str, candidate_name: str = "Candidate"
    ) -> Dict[str, Any]:
        score_guidelines = """
ATS SCORING RULES:
- 90–100: Excellent fit — nearly all required skills and tools match.
- 75–89: Good fit — strong alignment but 1–2 major gaps.
- 60–74: Average fit — partial match with several missing skills.
- 40–59: Below Average — limited overlap.
- 0–39: Poor fit — resume unrelated or lacks essentials.
"""

        prompt = f"""
You are an expert ATS and recruitment analyst.

{score_guidelines}

JOB ROLE: {job_role}
JOB DESCRIPTION:
{job_description}

CANDIDATE NAME: {candidate_name}
RESUME TEXT:
{cv_text}

TASK:
1. Identify required skills and experience from the job description.
2. Compare with the resume.
3. Apply the ATS SCORING RULES strictly.
4. Produce the final output ONLY in valid JSON (no text explanation).

OUTPUT FORMAT SCHEMA (do not copy numeric values):
{{
  "summary": "<6–8 sentence evaluation>",
  "ats_score": <integer between 0 and 100>,
  "grade": "<Excellent|Good|Average|Below Average|Poor>",
  "suggestions": [
    {{
      "missing_skill": "<specific skill or qualification>",
      "suggested_action": "<concrete next step>",
      "priority": "<high|medium|low>"
    }}
  ]
}}

INTERNAL SEED: {datetime.now().timestamp()}
"""
        try:
            response = self.model.generate_content(prompt)
            
            # Handle response properly
            if hasattr(response, 'text'):
                result_text = response.text.strip()
            elif hasattr(response, 'parts'):
                result_text = ''.join(part.text for part in response.parts if hasattr(part, 'text')).strip()
            else:
                result_text = str(response).strip()

            # Extract JSON from markdown code blocks
            json_match = re.search(r'```json\s*(.*?)\s*```', result_text, re.DOTALL)
            if json_match:
                result_text = json_match.group(1).strip()
            
            # Parse JSON
            parsed_data = json.loads(result_text)
            
            # Validate structure
            if not isinstance(parsed_data, dict):
                raise ValueError("Response is not a valid JSON object")

            # Normalize score
            score = max(0, min(100, int(parsed_data.get("ats_score", 50))))
            parsed_data["ats_score"] = score

            # Assign grade based on score
            if score >= 90:
                parsed_data["grade"] = "Excellent"
            elif score >= 75:
                parsed_data["grade"] = "Good"
            elif score >= 60:
                parsed_data["grade"] = "Average"
            elif score >= 40:
                parsed_data["grade"] = "Below Average"
            else:
                parsed_data["grade"] = "Poor"

            # Ensure suggestions is a list
            if not isinstance(parsed_data.get("suggestions"), list):
                parsed_data["suggestions"] = []

            logger.info(f"Comprehensive analysis completed (ATS Score: {score})")
            return parsed_data

        except Exception as e:
            logger.error(f"Comprehensive analysis failed: {e}")
            logger.error(f"Response type: {type(response) if 'response' in locals() else 'undefined'}")
            return {
                "summary": "Analysis failed due to parsing issue.",
                "ats_score": 50,
                "grade": "Average",
                "suggestions": []
            }

# ==================== FASTAPI APP ====================
app = FastAPI(
    title="AI-Powered Resume Screening & ATS Scoring System",
    description="Automated resume screening with deterministic scoring using Google Gemini AI",
    version="2.2.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS configuration - update with your frontend domain in production
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Change to specific domains in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

document_extractor = DocumentExtractor()
text_cleaner = TextCleaner()
ai_service = AIService()

# ==================== EXCEPTION HANDLERS ====================
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content=ErrorResponse(error=f"HTTP {exc.status_code}", detail=str(exc.detail)).dict()
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Unexpected error: {exc}")
    return JSONResponse(
        status_code=500,
        content=ErrorResponse(error="Internal Server Error", detail=str(exc)).dict()
    )

# ==================== ENDPOINTS ====================
@app.get("/", response_model=HealthCheck)
async def health_check():
    """Health check endpoint for monitoring"""
    return HealthCheck(
        status="healthy",
        message="AI Resume Screening & ATS Scoring System v2.3 is operational (supports PDF, DOCX, DOC)",
        timestamp=datetime.now().isoformat()
    )


@app.post("/api/v1/complete-analysis", response_model=CompleteAnalysisResponse)
async def complete_analysis(
    job_role: str = Form(...),
    job_description: str = Form(...),
    resume: UploadFile = File(...)
):
    """Complete resume analysis endpoint - supports PDF, DOCX, and DOC files"""
    
    # Validate file extension
    filename = resume.filename.lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx") or filename.endswith(".doc")):
        raise HTTPException(
            status_code=400, 
            detail="Only PDF (.pdf), Word Document (.docx), and Legacy Word (.doc) files are supported."
        )

    content = await resume.read()

    try:
        # Extract text based on file type
        raw_text = document_extractor.extract_text(content, resume.filename)
        cleaned_text = text_cleaner.clean_text(raw_text)
        
        logger.info(f"Processing resume for job role: {job_role}")

        # Extract candidate data
        candidate_data = ai_service.extract_candidate_data(cleaned_text)
        
        # Safely extract candidate details
        candidate_details_dict = candidate_data.get("candidate_details", {})
        if not isinstance(candidate_details_dict, dict):
            logger.warning(f"candidate_details is not a dict: {type(candidate_details_dict)}")
            candidate_details_dict = {}
        
        # Extract work experience
        work_exp_list = candidate_details_dict.get("work_experience", [])
        if not isinstance(work_exp_list, list):
            logger.warning(f"work_experience is not a list: {type(work_exp_list)}")
            work_exp_list = []
        
        work_experience_objs = []
        for exp in work_exp_list:
            try:
                if isinstance(exp, dict):
                    work_experience_objs.append(WorkExperience(**exp))
            except Exception as exp_error:
                logger.warning(f"Failed to parse work experience: {exp_error}")
                continue
        
        # Extract education
        education_list = candidate_details_dict.get("education", [])
        if not isinstance(education_list, list):
            logger.warning(f"education is not a list: {type(education_list)}")
            education_list = []
        
        education_objs = []
        for edu in education_list:
            try:
                if isinstance(edu, dict):
                    education_objs.append(Education(**edu))
            except Exception as edu_error:
                logger.warning(f"Failed to parse education: {edu_error}")
                continue
        
        # Create candidate details with all fields
        candidate_details = CandidateDetails(
            name=candidate_details_dict.get("name"),
            email=candidate_details_dict.get("email"),
            phone=candidate_details_dict.get("phone"),
            location=candidate_details_dict.get("location"),
            work_experience=work_experience_objs,
            education=education_objs,
            total_experience_years=candidate_details_dict.get("total_experience_years")
        )
        
        # Safely extract keywords
        keywords_data = candidate_data.get("keywords", {})
        if not isinstance(keywords_data, dict):
            logger.warning(f"keywords is not a dict: {type(keywords_data)}")
            keywords_data = {"keywords": []}
        
        keywords_list = keywords_data.get("keywords", [])
        if not isinstance(keywords_list, list):
            logger.warning(f"keywords list is not a list: {type(keywords_list)}")
            keywords_list = []
        
        # Convert keywords to KeywordItem objects
        keyword_objs = []
        for kw in keywords_list:
            try:
                if isinstance(kw, dict):
                    keyword_objs.append(KeywordItem(**kw))
                elif isinstance(kw, str):
                    keyword_objs.append(KeywordItem(keyword=kw, relevance_score=70))
                else:
                    logger.warning(f"Unexpected keyword type: {type(kw)}")
            except Exception as kw_error:
                logger.warning(f"Failed to parse keyword: {kw_error}")
                continue

        resume_data = ResumeData(
            job_role=job_role,
            job_description=job_description,
            candidate_details=candidate_details,
            keywords=Keywords(keywords=keyword_objs)
        )

        # Perform comprehensive analysis
        screening_data = ai_service.analyze_resume_comprehensive(
            job_role, job_description, cleaned_text, candidate_details.name or "Candidate"
        )
        
        # Safely extract suggestions
        suggestions_list = screening_data.get("suggestions", [])
        if not isinstance(suggestions_list, list):
            logger.warning(f"suggestions is not a list: {type(suggestions_list)}")
            suggestions_list = []
        
        # Convert suggestions to Suggestion objects
        suggestion_objs = []
        for s in suggestions_list:
            try:
                if isinstance(s, dict):
                    suggestion_objs.append(Suggestion(**s))
                else:
                    logger.warning(f"Unexpected suggestion type: {type(s)}")
            except Exception as s_error:
                logger.warning(f"Failed to parse suggestion: {s_error}")
                continue

        screening = ScreeningAnalysis(
            summary=screening_data.get("summary", "Analysis completed."),
            ats_score=screening_data.get("ats_score", 50),
            grade=screening_data.get("grade", "Average"),
            suggestions=suggestion_objs
        )

        return CompleteAnalysisResponse(
            success=True,
            resume_data=resume_data,
            screening=screening,
            analyzed_at=datetime.now().isoformat()
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis failed with error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Analysis failed: {str(e)}"
        )
    
@app.post("/api/v1/generate-report")
async def generate_report(analysis_request: ReportRequest = Body(...)):
    """
    Generate a professional text report based on AI analysis JSON.
    The frontend can use this text to render and download a PDF.
    """
    try:
        # Extract data from request
        data = analysis_request.analysis_data
        candidate = data.resume_data.candidate_details
        screening = data.screening
        keywords = [kw.keyword for kw in data.resume_data.keywords.keywords]
        suggestions = screening.suggestions

        # Prepare structured prompt for LLM
        report_prompt = f"""
You are an expert technical report writer specializing in HR analytics.

TASK: Generate a professional textual report summarizing the candidate's screening results.
The report will be used for a downloadable PDF document.

DATA INPUT:
- Candidate Name: {candidate.name or "N/A"}
- Email: {candidate.email or "N/A"}
- Phone: {candidate.phone or "N/A"}
- Location: {candidate.location or "N/A"}
- Total Experience: {candidate.total_experience_years or "N/A"} years
- Job Role Applied: {data.resume_data.job_role}
- ATS Score: {screening.ats_score}
- Grade: {screening.grade}
- Keywords Matched: {', '.join(keywords) if keywords else "N/A"}
- Screening Summary: {screening.summary}

SUGGESTIONS:
{chr(10).join([f"- {s.missing_skill}: {s.suggested_action} (Priority: {s.priority})" for s in suggestions]) or "No specific suggestions."}

FORMAT REQUIREMENTS:
1. Write in formal business English.
2. Structure it into clear sections: Overview, Candidate Summary, Match Evaluation, Key Skills, Recommendations, and Final Assessment.
3. Avoid using bullet points for entire report—mix narrative and concise lists.
4. End with a closing statement indicating overall suitability for the role.
"""

        # Use the existing AI service
        ai = AIService()
        response = ai.model.generate_content(report_prompt)

        # Extract text safely
        if hasattr(response, 'text'):
            report_text = response.text.strip()
        elif hasattr(response, 'parts'):
            report_text = ''.join(part.text for part in response.parts if hasattr(part, 'text')).strip()
        else:
            report_text = str(response).strip()

        if not report_text:
            raise ValueError("Empty response from LLM while generating report")

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "generated_report_text": report_text,
                "generated_at": datetime.now().isoformat()
            }
        )

    except Exception as e:
        logger.error(f"Report generation failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Report generation failed: {str(e)}"
        )

# ==================== STARTUP & SHUTDOWN ====================
@app.on_event("startup")
async def startup_event():
    logger.info("=" * 60)
    logger.info("AI Resume Screening System v2.3 Starting...")
    logger.info("Architecture: Deterministic Two-Prompt System")
    logger.info("  - Prompt 1: Candidate Extraction (with Experience & Education)")
    logger.info("  - Prompt 2: ATS Scoring with Rubric Logic")
    logger.info("Supported Formats: PDF, DOCX, DOC")
    logger.info(f"Environment: {os.getenv('ENVIRONMENT', 'production')}")
    logger.info("=" * 60)

@app.on_event("shutdown")
async def shutdown_event():
    logger.info("Shutting down Resume Screening System...")


